from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parent / "Bao_cao_nghiep_vu_va_huong_dan_SkinAI.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "198754"
RED = "C9364A"
GOLD = "7A5A00"
WHITE = "FFFFFF"
GRAY = "666666"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_style(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SkinAI | Báo cáo nghiệp vụ và hướng dẫn sử dụng")
    set_font(r, size=8.5, color=GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SkinAI | Trang ")
    set_font(r, size=8.5, color=GRAY)
    add_page_field(p)
    for run in p.runs:
        set_font(run, size=8.5, color=GRAY)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(88)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=27, color=INK, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        r = p.add_run(subtitle)
        set_font(r, size=14, color=DARK_BLUE)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text.upper())
    set_font(r, size=10, color=GREEN, bold=True)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, title, text, color=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        shade(cells[0], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        r = p.add_run(label)
        set_font(r, bold=True, color=DARK_BLUE, size=10)
        p = cells[1].paragraphs[0]
        r = p.add_run(value)
        set_font(r, size=10)
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row)):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (0, len(row) - 1) and len(text) < 26 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_function_block(doc, name, actor, entry, steps, results, notes=None):
    h = doc.add_heading(name, level=2)
    set_keep_with_next(h)
    add_kv_table(doc, [
        ("Người thực hiện", actor),
        ("Điểm truy cập", entry),
        ("Kết quả chính", results),
    ])
    h3 = doc.add_heading("Luồng thao tác", level=3)
    set_keep_with_next(h3)
    for step in steps:
        add_step(doc, step)
    if notes:
        add_callout(doc, "Lưu ý nghiệp vụ", notes)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    set_style(doc)
    add_header_footer(section)

    # Cover
    add_kicker(doc, "Tài liệu vận hành hệ thống")
    add_title(doc, "BÁO CÁO NGHIỆP VỤ\nVÀ HƯỚNG DẪN SỬ DỤNG", "Hệ thống SkinAI")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Phạm vi: các chức năng đã tích hợp trong phiên bản SkinAI hiện tại")
    set_font(r, size=11, color=GRAY, italic=True)
    add_kv_table(doc, [
        ("Ngày lập", "20/07/2026"),
        ("Đối tượng sử dụng", "Khách/Patient, Bác sĩ, Quản trị viên/Admin"),
        ("Mục đích", "Hướng dẫn demo, kiểm thử nghiệp vụ và vận hành hệ thống"),
        ("Phiên bản", "Tài liệu nội bộ - dựa trên mã nguồn SkinAI hiện tại"),
    ])
    doc.add_page_break()

    # Purpose and map
    doc.add_heading("1. Phạm vi và cách sử dụng tài liệu", level=1)
    add_body(doc, "Tài liệu này mô tả các luồng nghiệp vụ đã được bổ sung hoặc tích hợp cho SkinAI. Mỗi mục nêu rõ người dùng nào thực hiện, vào từ đâu, nút hoặc thao tác nào cần chọn, hệ thống xử lý gì và trạng thái nào được hiển thị sau đó.")
    add_callout(doc, "Cách đọc", "Các đường dẫn được ghi theo dạng /patient/... hoặc /admin/.... Khi chạy cục bộ, hệ thống có context path /SkinAI; ví dụ /SkinAI/patient/booking.")
    doc.add_heading("1.1 Vai trò và phạm vi quyền", level=2)
    add_matrix(doc,
               ["Vai trò", "Phạm vi chính", "Không thuộc phạm vi"],
               [
                   ["Khách", "Xem trang chủ, định vị phòng khám, đăng nhập", "Không xem dữ liệu bệnh nhân"],
                   ["Patient", "Hồ sơ, người thân, đặt lịch, thanh toán, bệnh án, đánh giá, báo lỗi", "Không xử lý báo lỗi của người khác"],
                   ["Doctor", "Xem menu bác sĩ và gửi báo lỗi/hỗ trợ", "Không quản trị issue report"],
                   ["Admin", "Quản lý đánh giá, xử lý issue report, xem thống kê/quản trị", "Không đặt lịch thay bệnh nhân"],
               ], [1500, 4500, 3360])
    doc.add_heading("1.2 Bản đồ điều hướng", level=2)
    add_matrix(doc,
               ["Nhu cầu", "Trang/nút cần mở", "Kết quả"],
               [
                   ["Tìm phòng khám", "Thanh menu: Phòng khám", "Mở /clinics và bản đồ gần bạn"],
                   ["Đặt lịch", "Menu tài khoản: Lịch hẹn hoặc nút đặt lịch", "Mở /patient/booking"],
                   ["Thanh toán", "Tự chuyển sau khi đặt lịch hoặc từ lịch hẹn", "Mở hóa đơn tại /patient/payment"],
                   ["Bệnh án", "Menu tài khoản: Hồ sơ bệnh án", "Mở /patient/medical-records"],
                   ["Lịch sử AI", "Trang chủ: Lịch sử chẩn đoán AI", "Mở /patient/reports"],
                   ["Báo lỗi", "Hồ sơ: Báo lỗi/Hỗ trợ", "Mở /patient/issue-report"],
               ], [2200, 3400, 3760])

    # Profile and family
    doc.add_heading("2. Hồ sơ cá nhân và thành viên gia đình", level=1)
    add_function_block(doc, "2.1 Xem và cập nhật hồ sơ cá nhân", "Patient", "Menu tên tài khoản -> Hồ sơ; đường dẫn /patient/profile", [
        "Di chuột hoặc nhấn vào tên tài khoản ở góc phải để mở menu tài khoản.",
        "Chọn Hồ sơ.",
        "Thay đổi các trường thông tin cá nhân được cho phép.",
        "Nhấn Lưu thay đổi. Hệ thống kiểm tra dữ liệu và cập nhật hồ sơ hiện tại.",
    ], "Thông tin được lưu trên hồ sơ Patient; các menu liên quan như Lịch hẹn, Hóa đơn, Hồ sơ bệnh án và Báo lỗi/Hỗ trợ được hiển thị từ khu vực tài khoản.")
    add_function_block(doc, "2.2 Thêm thành viên gia đình", "Patient", "Hồ sơ -> Thành viên gia đình -> Thêm thành viên gia đình", [
        "Tại khu vực Thành viên gia đình, khi chưa có dữ liệu, nhấn Thêm thành viên gia đình ở giữa khung.",
        "Nhập họ tên, ngày sinh, quan hệ, điện thoại và các thông tin địa chỉ cần thiết.",
        "Chọn quan hệ riêng biệt: Bố, Mẹ, Vợ/Chồng, Con, Anh, Chị, Em trai, Em gái, Ông/Bà hoặc Khác.",
        "Nhấn nút thêm/lưu. Hệ thống kiểm tra ngày sinh không ở tương lai và dữ liệu bắt buộc.",
        "Sau khi thành công, quay lại hồ sơ; danh sách hiển thị hai cột Người thân và Thông tin cá nhân, có nút Xem.",
    ], "Tạo một bản ghi family_members gắn với tài khoản chủ. Người thân không có tài khoản đăng nhập riêng.", "Giao diện không yêu cầu chọn giới tính. Backend dùng giá trị trung tính để tương thích với database cũ có cột gender bắt buộc.")
    add_function_block(doc, "2.3 Xem chi tiết người thân", "Patient", "Hồ sơ -> bảng Thành viên gia đình -> Xem", [
        "Nhấn Xem ở dòng người thân cần kiểm tra.",
        "Hệ thống chỉ tìm bản ghi có owner_user_id trùng tài khoản đang đăng nhập.",
        "Trang chi tiết hiển thị dữ liệu người thân đã đăng ký.",
    ], "Không thể mở thông tin người thân của một tài khoản khác.")

    # Booking
    doc.add_heading("3. Tìm bác sĩ và đặt lịch khám", level=1)
    add_function_block(doc, "3.1 Tìm bác sĩ theo điều kiện", "Patient", "Trang /patient/booking -> khung Tìm kiếm bác sĩ", [
        "Nhập tùy chọn tên bác sĩ, từ ngày, đến ngày, chuyên khoa hoặc ca làm việc.",
        "Nhấn nút kính lúp màu xanh để tìm.",
        "Hệ thống chỉ trả về bác sĩ đang hoạt động có ca phù hợp, còn chỗ, trong khoảng ngày/ca đã chọn.",
        "Kết quả hiển thị số lượng bác sĩ và thông tin họ tên, chuyên khoa, phòng khám.",
        "Nhấn Xóa bộ lọc để trở về trang đặt lịch không có điều kiện tìm kiếm.",
    ], "Mỗi doctor.id chỉ xuất hiện một lần dù có nhiều ca làm việc; nhiều bác sĩ khác nhau vẫn hiển thị đầy đủ.", "Ngày được lọc theo doctor_schedules.schedule_date; ca phải is_available = 1 và booked_count nhỏ hơn max_patients.")
    add_function_block(doc, "3.2 Chọn người khám", "Patient", "Trang /patient/booking -> mục Người khám", [
        "Mặc định chọn Tôi.",
        "Nếu đã thêm người thân, mở danh sách và chọn tên người thân mong muốn.",
        "Hệ thống gửi SELF hoặc FAMILY:<id> khi đặt lịch.",
        "Backend kiểm tra người thân đó thật sự thuộc tài khoản hiện tại trước khi tạo lịch.",
    ], "Lịch được gắn family_member_id khi đặt hộ; nếu chọn Tôi thì family_member_id để trống.")
    add_function_block(doc, "3.3 Chọn phòng khám, bác sĩ, ngày và ca", "Patient", "Các mục chọn trong biểu mẫu đặt lịch", [
        "Chọn phòng khám. Danh sách bác sĩ của phòng khám được tải lại.",
        "Chọn bác sĩ. Hệ thống tải các ngày có ca còn chỗ của bác sĩ đó.",
        "Chọn ngày. Nếu chưa chọn bác sĩ, hệ thống tải các bác sĩ còn ca trong ngày tại phòng khám; nếu đã chọn bác sĩ, hệ thống tải ca của bác sĩ đó.",
        "Nhấn ca Sáng, Chiều hoặc Tối. Hệ thống lưu slotId và tự tạo appointmentTime.",
        "Xem lại phần tóm tắt lịch trước khi gửi.",
    ], "Chỉ ca còn chỗ mới có thể được chọn.")
    add_function_block(doc, "3.4 Khai báo dị ứng", "Patient", "Trang /patient/booking -> Khai báo dị ứng (tùy chọn)", [
        "Nhập thông tin như: dị ứng penicillin, dị ứng hải sản hoặc phản ứng khi dùng mỹ phẩm.",
        "Có thể để trống.",
        "Khi chọn Tôi là người khám, thông tin được lưu lại vào hồ sơ Patient để lần sau tự hiển thị.",
        "Khi chọn người thân, thông tin dị ứng của hồ sơ chủ không bị ghi đè.",
    ], "Nội dung giới hạn 1.000 ký tự.")
    add_function_block(doc, "3.5 Xác nhận đặt lịch", "Patient", "Trang /patient/booking -> nút đặt lịch/xác nhận", [
        "Kiểm tra đủ clinicId, doctorId, slotId, thời gian lịch và requestId.",
        "Hệ thống chặn lịch chưa hoàn thành của đúng người được khám để tránh đặt trùng.",
        "Tạo lịch trong transaction và dùng requestId để chống gửi trùng yêu cầu.",
        "Sau khi thành công, hệ thống tự chuyển sang /patient/payment?action=create&appointmentId=... để bắt buộc chọn phương thức thanh toán.",
    ], "Lịch mới có trạng thái cuộc hẹn CONFIRMED và trạng thái tham gia NOT_VISITED.")
    add_callout(doc, "Quy tắc dọn dữ liệu lịch hẹn", "Khi tạo lịch thành công, hệ thống xóa các lịch đã hủy của bệnh nhân đó, sau đó chỉ giữ tối đa 5 lịch mới nhất. Các payment, invoice và feedback phụ thuộc của lịch bị xóa sẽ được dọn trước.", "FFF8E8")

    # Payments
    doc.add_heading("4. Hóa đơn và thanh toán", level=1)
    add_body(doc, "Ngay sau khi đặt lịch, hệ thống tạo hoặc tái sử dụng hóa đơn phí khám và chuyển người dùng sang trang thanh toán. Người dùng phải chọn một trong hai phương thức để tiếp tục.")
    add_function_block(doc, "4.1 Xem hóa đơn", "Patient", "Sau khi đặt lịch hoặc từ Lịch hẹn -> biểu tượng thanh toán", [
        "Hệ thống tạo/tìm hóa đơn theo appointmentId để không tạo trùng hóa đơn.",
        "Trang hiển thị thông tin lịch, tổng tiền, trạng thái hóa đơn và hai lựa chọn thanh toán.",
        "Bên dưới hiển thị cảnh báo: nếu thanh toán online thành công mà hủy lịch thì không hoàn tiền.",
    ], "Hóa đơn bắt đầu ở trạng thái UNPAID.")
    add_function_block(doc, "4.2 Gửi yêu cầu thanh toán tại quầy", "Patient", "Trang thanh toán -> nút Gửi yêu cầu tại quầy", [
        "Nhấn Gửi yêu cầu tại quầy.",
        "Hệ thống tạo payment có payment_method = CASH và status = PENDING.",
        "Quay lại danh sách lịch hẹn.",
        "Cột Trạng thái thanh toán hiển thị màu vàng: Đang chờ thanh toán.",
        "Người dùng vẫn có thể mở lại hóa đơn để đổi sang thanh toán online hoặc hủy lịch nếu nghiệp vụ cho phép.",
    ], "Gửi yêu cầu tại quầy không làm lịch hẹn thành Hoàn thành và không đổi trạng thái tham gia thành Đã khám.")
    add_function_block(doc, "4.3 Thanh toán online qua VNPay", "Patient", "Trang thanh toán -> nút Thanh toán online", [
        "Nhấn Thanh toán online.",
        "Form gửi POST đến Payment API Node.js: /api/invoices/{invoiceId}/payments/vnpay.",
        "Payment API kiểm tra hóa đơn UNPAID, tạo payment VNPAY/PENDING, tạo txn_ref và ký URL VNPay.",
        "Trình duyệt được chuyển đến cổng VNPay để chọn ngân hàng/QR/thẻ.",
        "VNPay gọi IPN/Return về Payment API. API kiểm tra checksum, txn_ref, số tiền và trạng thái giao dịch.",
        "Nếu hợp lệ và thành công, payment chuyển SUCCESS, invoice chuyển PAID và người dùng được quay về luồng kết quả/lịch hẹn.",
    ], "Thanh toán online có thời hạn 3 phút 30 giây tính từ lúc tạo giao dịch.", "VNPay là nguồn xác nhận thanh toán. Không được đánh dấu PAID chỉ dựa vào thao tác người dùng quay về trang kết quả.")
    add_function_block(doc, "4.4 Hết hạn thanh toán", "Payment API Node.js", "Bộ quét expirePendingPayments chạy nền", [
        "Payment PENDING có expires_at đã qua được đổi thành EXPIRED.",
        "Invoice UNPAID tương ứng chuyển CANCELLED.",
        "Appointment tương ứng chuyển CANCELLED và attendance_status chuyển CANCELLED.",
        "Nếu VNPay gửi callback muộn nhưng thời điểm thanh toán đã ký vẫn nằm trước deadline, giao dịch vẫn có thể được xác nhận và lịch được khôi phục về CONFIRMED.",
    ], "Bộ quét hiện chạy mỗi 5 phút; do đó thời hạn nghiệp vụ là 3 phút 30 giây nhưng thời điểm dọn dữ liệu thực tế có thể trễ hơn theo chu kỳ quét.")

    # Appointments and statuses
    doc.add_heading("5. Quản lý lịch hẹn, thanh toán và tham gia", level=1)
    doc.add_heading("5.1 Ý nghĩa hai cột trạng thái", level=2)
    add_matrix(doc,
               ["Cột", "Giá trị", "Ý nghĩa hiển thị"],
               [
                   ["Trạng thái thanh toán", "UNPAID", "Chưa tạo/gửi yêu cầu thanh toán thành công"],
                   ["Trạng thái thanh toán", "PENDING", "Đang chờ thanh toán, ví dụ yêu cầu tại quầy"],
                   ["Trạng thái thanh toán", "PAID", "Đã thanh toán thành công"],
                   ["Trạng thái thanh toán", "CANCELLED", "Hóa đơn đã hủy"],
                   ["Trạng thái tham gia", "NOT_VISITED", "Chưa khám"],
                   ["Trạng thái tham gia", "VISITED", "Đã khám"],
                   ["Trạng thái tham gia", "NO_SHOW", "Không có mặt sau khi quá giờ"],
                   ["Trạng thái tham gia", "CANCELLED", "Đã hủy lịch"],
               ], [2200, 2100, 5060])
    add_function_block(doc, "5.2 Xem lịch hẹn", "Patient", "Menu tài khoản -> Lịch hẹn; đường dẫn /patient/appointments", [
        "Mở Lịch hẹn.",
        "Hệ thống tự kiểm tra các lịch đã quá thời gian mà chưa tham gia.",
        "Hiển thị mã lịch, phòng khám, ngày giờ, trạng thái thanh toán, trạng thái tham gia, mục đích, ghi chú và thao tác.",
        "Các lịch thanh toán tại quầy đang chờ sẽ có trạng thái thanh toán màu vàng.",
    ], "Tại thời điểm tải trang, lịch CONFIRMED/CREATED đã quá giờ và vẫn NOT_VISITED sẽ được đổi thành NO_SHOW.")
    add_function_block(doc, "5.3 Hủy lịch hẹn", "Patient", "Lịch hẹn -> nút Xóa/Hủy -> xác nhận trong hộp thoại", [
        "Nhấn nút hủy ở lịch đủ điều kiện.",
        "Hộp thoại hỏi: Liệu bạn có muốn hủy lịch hẹn? Người dùng chọn Giữ lịch hẹn hoặc Xác nhận hủy.",
        "Khi xác nhận, hệ thống cập nhật appointment = CANCELLED và attendance_status = CANCELLED.",
        "Payment PENDING liên quan chuyển FAILED; invoice UNPAID chuyển CANCELLED.",
        "Hệ thống tạo thông báo hủy lịch cho tài khoản bệnh nhân.",
    ], "Nếu invoice đã PAID online, hệ thống không hoàn tiền; cảnh báo này được hiển thị trên trang thanh toán.")
    add_function_block(doc, "5.4 Đánh giá sau khám", "Patient", "Lịch hẹn hoàn thành -> nút Đánh giá hoặc menu Đánh giá của tôi", [
        "Chỉ lịch có appointment.status = COMPLETED và attendance_status = VISITED mới hiện/cho phép đánh giá.",
        "Chọn số sao từ 1 đến 5, loại đánh giá và nội dung.",
        "Nhấn gửi. Hệ thống kiểm tra bệnh nhân sở hữu lịch và mỗi lịch chỉ có một đánh giá.",
        "Người dùng có thể mở danh sách đánh giá của mình để xem hoặc cập nhật đánh giá.",
    ], "Không thể tạo đánh giá cho lịch hủy, chưa khám hoặc không có mặt.")

    # Notifications
    doc.add_heading("6. Thông báo và email", level=1)
    add_function_block(doc, "6.1 Nhận email thanh toán thành công", "Hệ thống/Patient", "Tự động sau khi invoice chuyển PAID", [
        "Payment API cập nhật payment SUCCESS và invoice PAID sau IPN hợp lệ từ VNPay.",
        "NotificationScheduler của Java chạy nền, lần đầu sau 10 giây và sau đó mỗi 30 giây.",
        "NotificationService tìm các hóa đơn PAID chưa có notification thành công.",
        "Tạo notification với event key duy nhất để tránh gửi trùng.",
        "Gửi email qua cấu hình SMTP và đánh dấu email_sent hoặc email_failed.",
    ], "Email gồm mã hóa đơn, số tiền, phương thức, thời điểm thanh toán, nội dung hóa đơn, phòng khám, bác sĩ, thời gian hẹn và mã giao dịch nếu có.")
    add_function_block(doc, "6.2 Thông báo hủy lịch", "Hệ thống/Patient", "Tự động khi Patient xác nhận hủy lịch", [
        "Sau transaction hủy lịch thành công, AppointmentController gọi queueAppointmentCancelled.",
        "Hệ thống tạo thông báo trong danh sách thông báo của người dùng.",
        "Thông báo chứa phòng khám, thời gian hẹn và lưu ý không hoàn tiền cho thanh toán online.",
        "Người dùng mở biểu tượng chuông hoặc /patient/notifications để xem và đánh dấu đã đọc.",
    ], "Thông báo trong app và email dùng cơ chế gửi có hàng đợi để tránh gửi trùng.")
    add_function_block(doc, "6.3 Thông báo đổi bác sĩ", "Hệ thống", "Sẵn sàng để gọi từ nghiệp vụ đổi bác sĩ", [
        "Hàm queueDoctorChanged nhận userId, appointmentId, tên bác sĩ cũ/mới và mã sự kiện.",
        "Tạo notification loại DOCTOR_CHANGED có event key riêng.",
        "Người dùng mở thông báo để quay về lịch hẹn liên quan.",
    ], "Hàm đã có nhưng cần được gọi từ transaction đổi bác sĩ của phòng khám/admin để trở thành luồng tự động hoàn chỉnh.")

    # Records
    doc.add_heading("7. Hồ sơ bệnh án, đơn thuốc và lịch sử AI", level=1)
    add_function_block(doc, "7.1 Xem hồ sơ bệnh án của tôi hoặc người thân", "Patient", "Menu tài khoản -> Hồ sơ bệnh án; đường dẫn /patient/medical-records", [
        "Mở Hồ sơ bệnh án.",
        "Ở tiêu đề, chọn Tôi hoặc chọn một người thân trong danh sách.",
        "Có thể tìm theo chẩn đoán/triệu chứng, lọc ngày bắt đầu-kết thúc và chọn sắp xếp mới nhất/cũ nhất.",
        "Hệ thống lấy medical_reports, nối lịch hẹn, bác sĩ, phòng khám và family_members.",
        "Chỉ hồ sơ có patient_id thuộc tài khoản hiện tại và đúng family_member_id được hiển thị.",
    ], "Hồ sơ bệnh án do bác sĩ tạo khác với lịch sử chẩn đoán AI.")
    add_function_block(doc, "7.2 Xem chi tiết hồ sơ và đơn thuốc", "Patient", "Bảng hồ sơ bệnh án -> Xem hoặc nút Đơn thuốc", [
        "Nhấn Xem để mở chi tiết chẩn đoán của bác sĩ, tình trạng, kế hoạch điều trị và ngày tái khám.",
        "Nếu có đơn thuốc, nhấn nút Đơn thuốc.",
        "Modal hiện danh sách thuốc gồm Tên thuốc, Số lượng và Liều dùng.",
        "Nhấn dấu X hoặc nhấn ra ngoài modal để đóng.",
    ], "Đơn thuốc được đọc từ appointment_prescriptions theo appointment_id; ghi chú đơn thuốc cũ không hiển thị trong modal.")
    add_function_block(doc, "7.3 Lịch sử chẩn đoán AI", "Patient", "Trang chủ -> liên kết gạch chân Lịch sử chẩn đoán AI; đường dẫn /patient/reports", [
        "Từ trang chủ, nhấn Lịch sử chẩn đoán AI dưới nút Bắt đầu chẩn đoán.",
        "Hệ thống mở danh sách diagnosis_reports của Patient.",
        "Người dùng xem bệnh được AI phát hiện, độ tin cậy, mức độ rủi ro, ngày tạo và chi tiết kết quả.",
    ], "Trang này không còn được gọi là Hồ sơ bệnh án, vì dữ liệu ở đây là kết quả AI chứ không phải bệnh án do bác sĩ lập.")

    # Issue reports
    doc.add_heading("8. Báo lỗi / Hỗ trợ (Issue Report)", level=1)
    add_function_block(doc, "8.1 Patient hoặc Doctor gửi báo lỗi", "Patient, Doctor", "Patient: Hồ sơ -> Báo lỗi/Hỗ trợ hoặc menu tài khoản. Doctor: menu tài khoản -> Báo lỗi/Hỗ trợ", [
        "Mở form báo lỗi.",
        "Nhập Tiêu đề sự cố, chọn Loại sự cố, nhập Mô tả lỗi; ảnh minh họa là tùy chọn.",
        "Có thể chọn các loại Lịch hẹn, Thanh toán, Tài khoản, Lỗi hệ thống hoặc Khác.",
        "Nếu đính kèm ảnh, chỉ chấp nhận JPG, PNG hoặc WEBP và tối đa 5 MB.",
        "Nhấn Gửi báo lỗi.",
        "Hệ thống sinh mã ISS-XXXXXXXX, lưu issue_reports ở trạng thái PENDING và ghi audit log.",
    ], "Form nhắc người dùng không gửi mật khẩu, OTP hoặc thông tin thẻ ngân hàng.")
    add_function_block(doc, "8.2 Admin xem và lọc báo lỗi", "Admin", "Thanh menu Admin -> Issue Reports; đường dẫn /admin/issue-reports", [
        "Mở Issue Reports.",
        "Dùng bộ lọc trạng thái hoặc ô tìm theo mã báo cáo, tiêu đề, người báo.",
        "Danh sách được ưu tiên theo PENDING, sau đó IN_PROGRESS và các trạng thái còn lại; trong mỗi nhóm, báo cáo mới hơn đứng trước.",
        "Xem tiêu đề, loại lỗi, người gửi, thời điểm, ảnh kèm theo và phản hồi trước đó.",
    ], "Admin phải có role ADMIN; người dùng khác nhận 403 khi cố gửi thao tác quản trị.")
    add_function_block(doc, "8.3 Admin xử lý và phản hồi báo lỗi", "Admin", "Issue Reports -> chọn thao tác trạng thái", [
        "Báo cáo PENDING: nhấn Xác nhận để chuyển IN_PROGRESS, hoặc chọn kết quả chưa thể xác nhận để chuyển REJECTED.",
        "Báo cáo IN_PROGRESS: nhấn Hoàn tất xử lý để chuyển RESOLVED, hoặc chọn kết quả chưa thể xác nhận để chuyển REJECTED.",
        "Hệ thống lưu admin_response, handled_by_admin_id, updated_at và resolved_at khi RESOLVED.",
        "Hệ thống ghi audit log thao tác của Admin.",
        "Nếu người gửi có email, hệ thống gửi email thông báo trạng thái và nội dung phản hồi chuyên nghiệp.",
    ], "Không thể chuyển trạng thái tùy ý; controller kiểm tra chuyển trạng thái hợp lệ để tránh bỏ qua bước xử lý.")
    doc.add_heading("8.4 Bảng trạng thái Issue Report", level=2)
    add_matrix(doc,
               ["Trạng thái", "Ý nghĩa", "Thao tác hợp lệ tiếp theo"],
               [
                   ["PENDING", "Đã tiếp nhận, chưa đánh giá", "Xác nhận -> IN_PROGRESS; Chưa thể xác nhận -> REJECTED"],
                   ["IN_PROGRESS", "Đã xác nhận lỗi, đang xử lý", "Hoàn tất -> RESOLVED; Chưa thể xác nhận -> REJECTED"],
                   ["RESOLVED", "Đã xử lý xong", "Kết thúc luồng"],
                   ["REJECTED", "Chưa thể xác nhận sự cố theo thông tin nhận được", "Kết thúc luồng; người dùng có thể gửi báo cáo mới với thêm thông tin"],
               ], [2200, 3600, 3560])

    # Locator and feedback
    doc.add_heading("9. Định vị phòng khám", level=1)
    add_function_block(doc, "9.1 Tìm cơ sở y tế gần vị trí hiện tại", "Khách, Patient", "Thanh menu -> Phòng khám; đường dẫn /clinics", [
        "Nhấn Phòng khám trên thanh menu.",
        "Trình duyệt yêu cầu quyền truy cập vị trí.",
        "Nếu đồng ý, hệ thống tìm các hospital/clinic/doctor trong bán kính 10 km từ vị trí hiện tại qua API bản đồ/địa điểm.",
        "Bản đồ MapTiler hiển thị vị trí người dùng và các cơ sở tìm được; nhấn cơ sở để xem thông tin và bay đến vị trí đó.",
        "Nếu người dùng từ chối định vị hoặc API ngoài lỗi, hệ thống chuyển sang dữ liệu phòng khám dự phòng từ database qua /api/clinic-fallback.",
    ], "Chỉ phòng khám active có latitude/longitude hợp lệ mới xuất hiện trong dữ liệu dự phòng.")
    add_callout(doc, "Ghi chú kỹ thuật", "Trang /clinics đang tải assets/javascript/app.js. File locate.js không phải file đang được trang này sử dụng, vì vậy không dùng nó để đánh giá kết quả vận hành hiện tại.")

    # Feedback
    doc.add_heading("10. Đánh giá phòng khám", level=1)
    add_function_block(doc, "10.1 Patient tạo hoặc sửa đánh giá", "Patient", "Lịch hẹn đã khám xong -> Đánh giá; hoặc menu -> Đánh giá của tôi", [
        "Từ lịch hẹn đủ điều kiện, nhấn nút Đánh giá.",
        "Chọn rating từ 1 đến 5 sao, chọn Khen/Góp ý/Khiếu nại và nhập nội dung tối đa 1.000 ký tự.",
        "Nhấn gửi để tạo feedback.",
        "Mở Đánh giá của tôi để xem danh sách; có thể nhấn Sửa để cập nhật đánh giá của chính mình.",
    ], "Mỗi appointment chỉ có một feedback. Backend xác minh appointment thuộc Patient và đã COMPLETED/VISITED.")
    add_function_block(doc, "10.2 Admin quản lý đánh giá", "Admin", "Thanh menu Admin -> Quản lý đánh giá; đường dẫn /admin/feedback", [
        "Mở danh sách đánh giá.",
        "Lọc theo trạng thái, loại đánh giá hoặc từ khóa; hệ thống phân trang kết quả.",
        "Nhấn xem chi tiết để xem Patient, lịch hẹn và nội dung đánh giá.",
        "Admin có thể phản hồi và cập nhật trạng thái PENDING, PROCESSING hoặc COMPLETED.",
        "Hệ thống ghi audit log khi phản hồi hoặc thay đổi trạng thái.",
    ], "Đánh giá là nghiệp vụ chất lượng dịch vụ; không dùng thay cho Issue Report về lỗi kỹ thuật.")

    # Operational notes
    doc.add_heading("11. Điều kiện vận hành và cấu hình", level=1)
    doc.add_heading("11.1 Thành phần hệ thống", level=2)
    add_matrix(doc,
               ["Thành phần", "Trách nhiệm", "Điều kiện cần"],
               [
                   ["Java/Tomcat", "Giao diện JSP, controller, DAO, notification scheduler", "JDK cấu hình cho module SkinAI; kết nối SQL Server"],
                   ["SQL Server", "Lưu Patient, lịch, hóa đơn, payment, báo lỗi, người thân, bệnh án, đơn thuốc", "Schema và migration đã được chạy đúng"],
                   ["Node.js Payment API", "Tạo URL VNPay, nhận IPN/Return, xử lý hết hạn thanh toán", "Chạy đúng cổng/cấu hình URL; Node dependencies"],
                   ["VNPay", "Xác nhận giao dịch online", "TMN code, secret key, Return URL/IPN URL hợp lệ"],
                   ["SMTP", "Gửi email thanh toán/thông báo báo lỗi", "Biến môi trường email hợp lệ"],
                   ["Nginx (nếu dùng)", "Reverse proxy/định tuyến khi triển khai", "Proxy đúng Tomcat và Payment API"],
               ], [2100, 3900, 3360])
    doc.add_heading("11.2 Bảng dữ liệu nghiệp vụ cần có", level=2)
    add_matrix(doc,
               ["Nhóm", "Bảng/cột quan trọng", "Được dùng bởi"],
               [
                   ["Người thân", "family_members; appointments.family_member_id", "Đặt lịch hộ, phân tách bệnh án"],
                   ["Lịch và tham gia", "appointments.status; attendance_status", "Lịch hẹn, no-show, hủy"],
                   ["Thanh toán", "invoices; payments; expires_at", "VNPay, tại quầy, trạng thái thanh toán"],
                   ["Bệnh án", "medical_reports; appointment_prescriptions", "Hồ sơ bác sĩ và modal đơn thuốc"],
                   ["Thông báo", "notifications", "In-app notification và hàng đợi email"],
                   ["Báo lỗi", "issue_reports", "Patient/Doctor gửi; Admin xử lý"],
                   ["Đánh giá", "feedbacks", "Patient đánh giá; Admin phản hồi"],
               ], [1900, 4100, 3360])
    add_callout(doc, "Không tự chạy lại schema khi demo", "Các migration/schema chỉ nên chạy sau khi backup database và xác nhận môi trường. Khi trình bày hoặc kiểm thử chức năng, ưu tiên dùng dữ liệu mẫu đã chuẩn bị để tránh xóa/ghi đè lịch hẹn và hóa đơn.", "FFF8E8")
    doc.add_heading("11.3 Danh sách kiểm thử nhanh trước demo", level=2)
    for item in [
        "Đăng nhập bằng tài khoản Patient có hồ sơ Patient đầy đủ.",
        "Xác nhận có ít nhất một phòng khám, bác sĩ active và doctor_schedules còn chỗ ở ngày demo.",
        "Kiểm tra Payment API Node.js đang chạy và URL trong payment.jsp truy cập được.",
        "Kiểm tra cấu hình VNPay dùng đúng môi trường test/sandbox khi demo.",
        "Kiểm tra SMTP/.env nếu cần chứng minh email thanh toán hoặc email báo lỗi.",
        "Chuẩn bị một lịch COMPLETED/VISITED để demo đánh giá và một medical_report/appointment_prescription để demo bệnh án.",
        "Chuẩn bị một tài khoản Admin để demo trạng thái Issue Report và email phản hồi.",
    ]:
        add_bullet(doc, item)

    # Quick scripts
    doc.add_heading("12. Kịch bản demo đề xuất", level=1)
    doc.add_heading("12.1 Kịch bản Patient: đặt lịch và thanh toán", level=2)
    for step in [
        "Đăng nhập Patient -> mở Hồ sơ -> thêm người thân.",
        "Mở Đặt lịch -> chọn người khám là người thân -> chọn phòng khám, bác sĩ, ngày và ca.",
        "Gửi đặt lịch -> chứng minh hệ thống tự chuyển sang Thanh toán.",
        "Chọn Gửi yêu cầu tại quầy -> quay lại Lịch hẹn và chỉ ra trạng thái thanh toán màu vàng.",
        "Mở lại hóa đơn -> chọn Thanh toán online -> cho thấy URL được chuyển đến VNPay test.",
        "Sau khi thanh toán thành công, mở Thông báo/email và Lịch hẹn để kiểm tra PAID.",
    ]:
        add_step(doc, step)
    doc.add_heading("12.2 Kịch bản Patient: bệnh án và đơn thuốc", level=2)
    for step in [
        "Mở menu tài khoản -> Hồ sơ bệnh án.",
        "Chuyển giữa Tôi và người thân để chứng minh dữ liệu được tách theo người khám.",
        "Nhấn nút Đơn thuốc của một hồ sơ -> mở modal danh sách thuốc.",
        "Đóng modal bằng dấu X hoặc nhấn bên ngoài.",
        "Quay về Trang chủ -> nhấn Lịch sử chẩn đoán AI để cho thấy đây là luồng AI tách biệt.",
    ]:
        add_step(doc, step)
    doc.add_heading("12.3 Kịch bản báo lỗi", level=2)
    for step in [
        "Từ Hồ sơ, nhấn Báo lỗi/Hỗ trợ.",
        "Nhập tiêu đề, chọn Thanh toán, mô tả, đính kèm ảnh hợp lệ và gửi.",
        "Đăng nhập Admin -> mở Issue Reports -> tìm theo mã ISS vừa tạo.",
        "Nhấn Xác nhận để chuyển IN_PROGRESS; kiểm tra email nếu SMTP đã cấu hình.",
        "Nhấn Hoàn tất xử lý để chuyển RESOLVED; kiểm tra lịch sử/audit và email phản hồi.",
    ]:
        add_step(doc, step)

    # Doctor handoff and shared patient-record contract
    doc.add_heading("13. Luồng phía Bác sĩ và điểm bàn giao với Patient", level=1)
    add_callout(doc, "Phạm vi đã có trong source hiện tại", "Doctor hiện đã có Dashboard, xem chi tiết lịch, duyệt/từ chối lịch, hoàn tất lịch và quản lý lịch làm việc. Phần lập mới bệnh án (medical_reports) và thêm thuốc (appointment_prescriptions) chưa có controller/DAO ghi dữ liệu ở source hiện tại; đây là phần cần đồng nghiệp Doctor nối theo hợp đồng dữ liệu bên dưới.", "FFF8E8")

    doc.add_heading("13.1 Thứ tự thao tác Doctor", level=2)
    add_function_block(doc, "Bước 1 - Mở Dashboard và lọc lịch", "Doctor", "Đăng nhập Doctor -> /doctor/dashboard", [
        "Hệ thống lấy doctor profile từ user đang đăng nhập và chỉ tải các appointment có doctor_id của Doctor đó.",
        "Doctor xem tổng số lịch và lọc theo doctor_status: PENDING, ACCEPTED hoặc REJECTED.",
        "Nhấn một lịch để vào trang chi tiết /doctor/appointments/detail?id={appointmentId}.",
    ], "Dashboard là điểm bắt đầu để xử lý yêu cầu khám; doctor_status là trạng thái duyệt của Doctor, khác với trạng thái tổng của lịch và trạng thái thanh toán.")
    add_function_block(doc, "Bước 2 - Xem thông tin ca khám", "Doctor", "Dashboard -> chọn lịch -> Chi tiết lịch", [
        "Xem bệnh nhân/người được khám, phòng khám, ngày giờ, triệu chứng/ghi chú và các thông tin cần thiết trước khi duyệt.",
        "Nếu cần, Doctor nhập doctorNotes để lưu lý do duyệt hoặc từ chối.",
        "Mọi request chi tiết phải xác minh appointment thuộc doctor_id của Doctor đăng nhập trước khi hiển thị hoặc cập nhật.",
    ], "Kiểm tra quyền sở hữu lịch là yêu cầu bắt buộc khi hoàn thiện module Doctor; không chỉ dựa vào appointmentId trên URL.")
    add_function_block(doc, "Bước 3 - Duyệt hoặc từ chối lịch", "Doctor", "Chi tiết lịch -> Duyệt hoặc Từ chối", [
        "Duyệt: cập nhật appointments.doctor_status = ACCEPTED, lưu ghi chú và cập nhật trạng thái lịch tổng thành CONFIRMED.",
        "Từ chối: cập nhật appointments.doctor_status = REJECTED, lưu lý do để bộ phận điều phối/Patient có thể hiểu kết quả.",
        "Sau khi cập nhật, trình duyệt quay lại cùng trang chi tiết với cờ success=true hoặc error=true để hiển thị thông báo.",
    ], "Không dùng trạng thái thanh toán để biểu diễn việc Doctor nhận hay từ chối ca khám.")
    add_function_block(doc, "Bước 4 - Quản lý lịch làm việc", "Doctor", "Menu Doctor -> /doctor/schedule", [
        "Chọn tuần trước/sau và bật/tắt khả dụng cho ca Sáng, Chiều, Tối.",
        "Hệ thống đọc doctor_schedules theo doctor_id, ngày và ca; đồng thời đếm lịch không bị CANCELLED hoặc NO_SHOW.",
        "Khi Patient tìm Doctor/đặt lịch, các API đặt lịch phải dùng dữ liệu khả dụng này để không hiển thị ca hết chỗ.",
    ], "DoctorSchedule là nguồn dữ liệu chung cho phía Doctor và luồng tìm/đặt lịch của Patient.")
    add_function_block(doc, "Bước 5 - Khám, lập bệnh án và đơn thuốc", "Doctor", "Chi tiết lịch đã khám -> Lưu bệnh án và đơn thuốc", [
        "Doctor nhập triệu chứng chính, chẩn đoán, kế hoạch điều trị, ngày tái khám và trạng thái hồ sơ; lưu một bản ghi medical_reports gắn với appointment_id và doctor_id.",
        "Doctor thêm từng thuốc gồm tên thuốc, số lượng và liều dùng; mỗi dòng là một appointment_prescriptions gắn với cùng appointment_id.",
        "Nên lưu medical_reports và toàn bộ thuốc trong một transaction: hoặc cả bệnh án và thuốc cùng thành công, hoặc cùng rollback khi lỗi.",
        "Chỉ sau khi lưu bệnh án thành công mới chuyển appointment sang COMPLETED để Patient xem được thông tin hoàn chỉnh.",
    ], "Đây là luồng cần bổ sung cho module Doctor. Patient hiện đã sẵn phần đọc và hiển thị dữ liệu sau khi Doctor ghi đúng cấu trúc.")
    add_function_block(doc, "Bước 6 - Hoàn tất ca khám", "Doctor", "Chi tiết lịch -> Hoàn thành", [
        "Cập nhật appointments.status = COMPLETED; AppointmentDAO đồng thời quy đổi attendance_status = VISITED.",
        "Nếu có completionNotes, hệ thống lưu vào doctor_notes với doctor_status = ACCEPTED.",
        "InvoiceService hiện được gọi để tạo hóa đơn cho lịch hoàn thành; cần thống nhất với team Payment để không tạo trùng với hóa đơn đã được tạo ở thời điểm đặt lịch.",
        "Patient khi xem lịch sẽ thấy đã khám và có thể đủ điều kiện gửi đánh giá; bệnh án/đơn thuốc được đọc theo appointment_id.",
    ], "Hoàn thành ca khám là một trạng thái nghiệp vụ khác với PAID. Không được suy ra đã khám chỉ vì hóa đơn đã thanh toán.")

    doc.add_heading("13.2 Bảng prescription ở đâu và tương tác thế nào", level=2)
    add_matrix(doc,
               ["Đối tượng", "Bảng/khóa", "Vai trò trong luồng"],
               [
                   ["Lịch khám", "appointments.id", "Khóa gốc cho đúng một ca khám; mang patient_id, family_member_id, doctor_id, clinic_id và trạng thái."],
                   ["Bệnh án", "medical_reports.appointment_id", "Bệnh án do Doctor lập; liên kết một ca khám với chẩn đoán, điều trị, tái khám và Doctor."],
                   ["Đơn thuốc", "appointment_prescriptions.appointment_id", "Mỗi dòng là một thuốc: drug_name, quantity, dosage. Khóa ngoại trỏ appointments(id), ON DELETE CASCADE."],
                   ["Giao diện Patient", "MedicalReportDAO.findPrescriptionsByAppointmentId", "Khi có report, hệ thống đếm và đọc thuốc theo appointment_id; nút Đơn thuốc mở modal hiển thị Tên thuốc, Số lượng, Liều dùng."],
               ], [1800, 2700, 4860])
    add_body(doc, "Điểm quan trọng: prescription không gắn trực tiếp bằng patient_id. Nó gắn vào appointment_id; từ appointment hệ thống suy ra đúng patient_id và, nếu đặt hộ, đúng family_member_id. Vì vậy một chủ tài khoản chỉ xem được đơn thuốc của chính mình hoặc người thân do mình quản lý.")
    add_callout(doc, "Hợp đồng dữ liệu cho Doctor", "Khi thêm form kê đơn, Doctor chỉ cần ghi medical_reports và appointment_prescriptions với cùng appointment_id. Patient không cần sửa luồng đọc nếu dữ liệu đó hợp lệ. Không lưu thuốc vào diagnosis_reports vì đó là lịch sử chẩn đoán AI, không phải bệnh án do Doctor lập.", "EAF5EF")

    doc.add_heading("13.3 Redirect kèm hash anchor (#...) là gì", level=2)
    add_body(doc, "Redirect là phản hồi HTTP yêu cầu trình duyệt mở một URL khác, ví dụ sau khi xử lý form: /doctor/appointments/detail?id={appointmentId}&success=true. Query string (id, success) được gửi đến server để controller đọc và hiển thị dữ liệu/thông báo.")
    add_body(doc, "Hash anchor là phần sau dấu # của URL, ví dụ /doctor/appointments/detail?id={appointmentId}#medical-record. Hash không được gửi lên server. Sau khi trang đã tải, chính trình duyệt dùng hash để cuộn đến phần HTML có id=\"medical-record\". Tác dụng là đưa Doctor hoặc Patient thẳng tới vùng vừa thao tác, tránh phải tự cuộn trang.")
    add_matrix(doc,
               ["Dạng URL", "Server nhận được", "Tác dụng ở giao diện"],
               [
                   ["...?id=A1&success=true", "id=A1, success=true", "Controller tải ca A1; JSP hiện thông báo thành công."],
                   ["...?id=A1#prescription", "Chỉ id=A1", "Sau khi JSP tải xong, trình duyệt cuộn tới phần có id=prescription."],
                   ["...?id=A1#prescriptionModal", "Chỉ id=A1", "Chỉ hash thì không tự mở Bootstrap modal; cần JavaScript đọc window.location.hash rồi gọi modal.show()."],
               ], [2800, 2500, 4060])
    add_callout(doc, "Quy tắc sử dụng", "Dùng query string cho dữ liệu/ý nghĩa cần backend xử lý (id, action result, filter). Dùng #hash chỉ để hỗ trợ điều hướng trải nghiệm trên chính trang. Hash không phải cơ chế phân quyền, không thay thế kiểm tra Doctor/Patient sở hữu dữ liệu.", "F4F6F9")

    doc.add_heading("13.4 Những phần Doctor tương tác trong dự án", level=2)
    add_matrix(doc,
               ["Khu vực", "Doctor tác động", "Phía Patient/hệ thống nhận tác động"],
               [
                   ["DoctorDashboardController + AppointmentDAO", "Lọc và duyệt/từ chối doctor_status; xác nhận lịch.", "Lịch Patient chuyển CONFIRMED hoặc phản ánh kết quả xử lý Doctor."],
                   ["DoctorScheduleController + DoctorScheduleDAO", "Thiết lập khả dụng theo ngày/ca, theo dõi số lịch.", "Tìm Doctor/đặt lịch dùng sức chứa và ca còn khả dụng."],
                   ["DoctorAppointmentDetailController", "Hoàn tất lịch, lưu ghi chú và gọi InvoiceService theo source hiện tại.", "Trạng thái tham gia Patient đổi VISITED; luồng hóa đơn cần được team thống nhất để tránh tạo trùng."],
                   ["medical_reports + appointment_prescriptions", "Đây là điểm Doctor sẽ ghi bệnh án và đơn thuốc.", "MedicalRecordsController/MedicalReportDAO đang đọc, lọc theo Tôi/người thân và hiển thị modal đơn thuốc."],
                   ["appointments.family_member_id", "Doctor cần xem đây là người thực sự được khám, không mặc định người đặt lịch là người khám.", "Bệnh án, đơn thuốc và lịch sử phải thuộc đúng người thân đã được chọn."],
                   ["Issue Report", "Doctor có thể gửi báo lỗi/hỗ trợ như Patient.", "Admin tiếp nhận, đổi trạng thái và gửi email phản hồi cho người báo."],
               ], [2550, 3300, 3510])
    doc.add_heading("13.5 Checklist tích hợp an toàn cho đồng nghiệp Doctor", level=2)
    for item in [
        "Mọi GET/POST theo appointmentId phải kiểm tra appointment.doctor_id khớp Doctor từ session trước khi đọc hoặc cập nhật.",
        "Khóa form lập bệnh án/kê đơn nếu lịch chưa CONFIRMED hoặc đã CANCELLED/NO_SHOW; không cho tạo bệnh án cho sai ca.",
        "Bảo đảm một medical_report cho một appointment (nên có unique constraint hoặc kiểm tra trước khi insert).",
        "Dùng transaction khi lưu medical_report và danh sách appointment_prescriptions.",
        "Không ghi patient_id hoặc family_member_id vào đơn thuốc theo dữ liệu form; luôn suy từ appointment ở server.",
        "Sau khi hoàn tất, redirect về trang chi tiết với success=true và có thể thêm #medical-record để người dùng nhìn thấy phần vừa lưu.",
        "Thống nhất thời điểm tạo invoice với team Payment; hiện source đã có cả luồng tạo hóa đơn khi đặt lịch và lời gọi khi Doctor complete, cần chọn một nguồn tạo duy nhất.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("14. Giới hạn đã biết và hướng cải tiến", level=1)
    add_bullet(doc, "Luồng email bệnh nhân - bác sĩ riêng tư mới ở mức ý tưởng; chưa triển khai thành email ticket độc lập.")
    add_bullet(doc, "Thông báo đổi bác sĩ đã có hàm tạo notification nhưng chưa được nối vào thao tác đổi bác sĩ thực tế.")
    add_bullet(doc, "Bộ dọn giao dịch hết hạn chạy mỗi 5 phút; nếu cần phản hồi sát mốc 3 phút 30 giây hơn, có thể giảm chu kỳ quét sau khi đánh giá tải hệ thống.")
    add_bullet(doc, "Issue Report hiện phục vụ Patient và Doctor gửi báo lỗi; cần quy trình phân công/phân loại sâu hơn nếu số lượng báo cáo tăng.")
    add_bullet(doc, "Khi triển khai ngoài localhost, URL Payment API và Nginx cần thay bằng domain thực tế; không giữ localhost:3000 trong giao diện production.")

    doc.add_heading("15. Kết luận", level=1)
    add_body(doc, "SkinAI hiện có luồng liên kết từ hồ sơ - người thân - đặt lịch - thanh toán - lịch hẹn - thông báo - hồ sơ bệnh án, đồng thời bổ sung các kênh hỗ trợ chất lượng gồm đánh giá và báo lỗi. Các trạng thái được tách rõ giữa thanh toán và tham gia khám để người dùng, lễ tân và quản trị viên không nhầm lẫn nghiệp vụ.")
    add_callout(doc, "Thông điệp demo", "Hệ thống không chỉ đặt lịch; hệ thống quản lý đúng người được khám, điều kiện thanh toán, trạng thái tham gia, dữ liệu bệnh án/đơn thuốc và phản hồi sự cố theo từng vai trò.", "EAF5EF")

    doc.core_properties.title = "Báo cáo nghiệp vụ và hướng dẫn sử dụng SkinAI"
    doc.core_properties.subject = "Tài liệu nghiệp vụ, luồng thao tác và hướng dẫn sử dụng SkinAI"
    doc.core_properties.author = "SkinAI Project Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
